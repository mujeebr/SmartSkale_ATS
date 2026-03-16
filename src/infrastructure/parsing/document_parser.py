"""Concrete `DocumentParser` implementation for PDF, DOCX and text resumes."""

from __future__ import annotations

from io import BytesIO

import docx
from PyPDF2 import PdfReader

from ...core.parsing.base import DocumentParser, ParsedDocument


class DefaultDocumentParser(DocumentParser):
    """Parse raw file bytes into `ParsedDocument` instances."""

    def parse(self, filename: str, data: bytes) -> ParsedDocument:
        """Dispatch parsing based on filename extension."""
        lower = filename.lower()
        if lower.endswith(".docx"):
            text = self._from_docx_bytes(data)
        elif lower.endswith(".pdf"):
            text = self._from_pdf_bytes(data)
        else:
            text = self._from_txt_bytes(data)

        return ParsedDocument(filename=filename, text=text)

    def _from_docx_bytes(self, data: bytes) -> str:
        """Extract plain text from a DOCX file held in memory."""
        doc = docx.Document(BytesIO(data))
        return "\n".join([para.text for para in doc.paragraphs]).strip()

    def _from_pdf_bytes(self, data: bytes) -> str:
        """Extract plain text from a PDF file held in memory."""
        reader = PdfReader(BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted:
                parts.append(extracted)
        return "\n".join(parts).strip()

    def _from_txt_bytes(self, data: bytes) -> str:
        """Decode and normalise plain-text resume content."""
        return data.decode("utf-8", errors="replace").strip()
